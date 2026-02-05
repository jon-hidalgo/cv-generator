#!/usr/bin/env python3
from docx import Document

doc = Document('my_cv.docx')

# Find KEY ACHIEVEMENTS section
print("=== KEY ACHIEVEMENTS ===")
start = None
for i, para in enumerate(doc.paragraphs):
    if 'KEY ACHIEVEMENTS' in para.text:
        start = i
        break

if start:
    for i in range(start, min(start + 10, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        text = para.text[:60] if para.text else "(empty)"
        print(f"Para {i}: {repr(text)}")
        pPr = para._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
        if pPr is not None:
            numPr = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
            print(f"  Has numPr: {numPr is not None}")
        else:
            print("  pPr is None")

print("\n=== TECHNICAL STACK ===")
start = None
for i, para in enumerate(doc.paragraphs):
    if 'TECHNICAL STACK' in para.text:
        start = i
        break

if start:
    for i in range(start, min(start + 10, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        text = para.text[:60] if para.text else "(empty)"
        print(f"Para {i}: {repr(text)}")
        pPr = para._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
        if pPr is not None:
            numPr = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
            print(f"  Has numPr: {numPr is not None}")
        else:
            print("  pPr is None")
