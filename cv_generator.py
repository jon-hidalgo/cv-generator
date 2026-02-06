#!/usr/bin/env python3
"""
CV Template Generator

This script takes a CV template with placeholders and fills them with provided values.
Placeholders should be in the format: {{PLACEHOLDER_NAME}}

Usage:
    python cv_generator.py --template template.txt --output output.txt --data data.json
    python cv_generator.py --template template.txt --output output.txt --name "John Doe" --email "john@example.com"
"""

import argparse
import json
import sys
import re
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
import subprocess # Re-import subprocess for direct soffice call

def load_template(template_path):
    """Load the template file (supports .txt and .docx)."""
    template_path = Path(template_path)
    
    if not template_path.exists():
        print(f"Error: Template file '{template_path}' not found.", file=sys.stderr)
        sys.exit(1)
    
    try:
        if template_path.suffix.lower() == '.docx':
            return Document(template_path)
        else:
            with open(template_path, 'r', encoding='utf-8') as f:
                return f.read()
    except Exception as e:
        print(f"Error reading template: {e}", file=sys.stderr)
        sys.exit(1)


def load_data_from_json(json_path):
    """Load replacement data from JSON file."""
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except FileNotFoundError:
        print(f"Error: Data file '{json_path}' not found.", file=sys.stderr)
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"Error parsing JSON: {e}", file=sys.stderr)
        sys.exit(1)


def _replace_placeholder_in_runs(paragraph, placeholder, replacement_value):
    """
    Replaces a specific placeholder string within a paragraph's runs,
    preserving existing run formatting. This function intelligently handles
    placeholders that may span across multiple runs by reconstructing the
    affected runs while maintaining their original formatting.
    """
    segments = [] # List of (text_content, original_run_properties_dict)
    
    # Collect all existing run data (text and properties)
    for run in paragraph.runs:
        segments.append({
            'text': run.text,
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'name': run.font.name,
            'size': run.font.size,
            'color': run.font.color.rgb
        })

    # Prepare for replacement
    new_segments_data = []
    found_replacement = False

    for segment in segments:
        text_to_process = segment['text']
        
        if placeholder in text_to_process:
            found_replacement = True
            parts = text_to_process.split(placeholder)
            
            if parts[0]:
                new_segments_data.append(dict(segment, text=parts[0]))
            
            new_segments_data.append(dict(segment, text=replacement_value))
            
            if parts[1:]:
                new_segments_data.append(dict(segment, text="".join(parts[1:])))
        else:
            new_segments_data.append(segment)

    if not found_replacement:
        return # No replacement occurred

    # Clear all existing runs in the paragraph
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    
    # Add new runs based on the `new_segments_data`
    for new_segment in new_segments_data:
        if not new_segment['text']:
            continue
        
        # Parse markdown within the text content of the new_segment
        markdown_parts = parse_markdown_formatting(new_segment['text'])
        
        for part_text, is_markdown_bold in markdown_parts:
            run = paragraph.add_run(part_text)
            
            # Apply inherited template formatting
            run.bold = new_segment.get('bold', False)
            run.italic = new_segment.get('italic', False)
            run.underline = new_segment.get('underline', False)
            if new_segment.get('name'):
                run.font.name = new_segment['name']
            if new_segment.get('size'):
                run.font.size = new_segment['size']
            if new_segment.get('color'):
                run.font.color.rgb = new_segment['color']
            
            # Apply markdown bolding (overriding inherited bold if markdown says it's bold)
            if is_markdown_bold:
                run.bold = True
            
def parse_markdown_formatting(text):
    """
    Parse markdown-style formatting in text.
    Returns a list of tuples: (text, is_bold)
    
    Example: "**bold** *also bold* text" -> [("bold", True), (" ", False), ("also bold", True), (" text", False)]
    """
    parts = []
    # Pattern to match **text** or *text*, using a backreference `\1` to ensure closing markers match opening ones.
    pattern = r'(\*{1,2})(.+?)\1'
    last_end = 0
    
    for match in re.finditer(pattern, text):
        # Add text before the bold part
        if match.start() > last_end:
            parts.append((text[last_end:match.start()], False))
        
        # The actual matched text to be bolded is in group 2
        bold_text = match.group(2)
        parts.append((bold_text, True))
        last_end = match.end()
    
    # Add remaining text
    if last_end < len(text):
        parts.append((text[last_end:], False))
    
    return parts if parts else [(text, False)]


def set_paragraph_text_with_formatting(paragraph, text, run_font_properties=None):
    """Set text of a paragraph applying markdown formatting and preserving run properties."""
    # Clear existing runs
    for run in paragraph.runs:
        run.clear()
    
    # Apply new text with formatting
    parts = parse_markdown_formatting(text)
    for i, (part_text, is_bold) in enumerate(parts):
        run = paragraph.add_run(part_text)
        if run_font_properties:
            run.bold = run_font_properties.get('bold', False)
            run.italic = run_font_properties.get('italic', False)
            run.underline = run_font_properties.get('underline', False)
            if run_font_properties.get('name'):
                run.font.name = run_font_properties['name']
            if run_font_properties.get('size'):
                run.font.size = run_font_properties['size']
            if run_font_properties.get('color'):
                run.font.color.rgb = run_font_properties['color']
        run.bold = is_bold # Apply or override bold based on markdown


def add_bullet_style(paragraph_element):
    """Add bullet point style to a paragraph element."""
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    pPr = paragraph_element.find('w:pPr', namespaces=nsmap)
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        paragraph_element.insert(0, pPr)
    
    numPr = pPr.find('w:numPr', namespaces=nsmap)
    if numPr is None:
        numPr = OxmlElement('w:numPr')
        pPr.append(numPr)
    
    ilvl = numPr.find('w:ilvl', namespaces=nsmap)
    if ilvl is None:
        ilvl = OxmlElement('w:ilvl')
        ilvl.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '0')
        numPr.append(ilvl)
        
    numId = numPr.find('w:numId', namespaces=nsmap)
    if numId is None:
        numId = OxmlElement('w:numId')
        numId.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '1') # Assuming '1' is a bullet list
        numPr.append(numId)


def remove_trailing_empty_paragraphs(document):
    """
    Removes trailing empty paragraphs from the end of the document.
    This often prevents an unwanted blank page at the end of the generated docx.
    """
    paragraphs = document.paragraphs
    # Iterate from the end of the document backwards
    for i in range(len(paragraphs) - 1, -1, -1):
        paragraph = paragraphs[i]
        # Check if the paragraph is empty (no text and no runs or only contains whitespace)
        # Using _element.xml to check for more robust emptiness (e.g., contains only <w:br/>)
        # However, a simple text check is usually sufficient for user-created empty paragraphs
        if not paragraph.text.strip() and not paragraph.runs:
            # If it's the last paragraph and empty, remove it
            if paragraph._element.getparent() is not None:
                paragraph._element.getparent().remove(paragraph._element)
        else:
            # Found a non-empty paragraph, stop removing
            break

def fill_docx_template(doc, data):
    """Fill DOCX template placeholders with data values, supporting markdown formatting and repeating blocks."""
    
    # Process repeating blocks first
    
    all_paragraphs = list(doc.paragraphs)
    
    i = 0
    while i < len(all_paragraphs):
        para = all_paragraphs[i]
        
        match = re.search(r'{{#(\w+)}}', para.text)
        if not match:
            i += 1
            continue

        block_name = match.group(1)
        
        if not (block_name in data and isinstance(data[block_name], list)):
            i+= 1
            continue

        # Find the closing paragraph
        closing_idx = -1
        for j in range(i + 1, len(all_paragraphs)):
            if f'{{{{/{block_name}}}}}' in all_paragraphs[j].text:
                closing_idx = j
                break
        
        if closing_idx == -1:
            i += 1
            continue
            
        opening_para = all_paragraphs[i]
        template_paras = all_paragraphs[i+1:closing_idx]
        
        # Remove template paragraphs including opening and closing markers
        # Store a reference to the element before the opening paragraph for insertion
        anchor_element = opening_para._element.getprevious()
        
        for p_idx in range(i, closing_idx + 1):
            p = all_paragraphs[p_idx]
            p._element.getparent().remove(p._element)

        # Collect all generated paragraphs for this block
        generated_paragraph_elements = []
        for item_data in data[block_name]:
            for template_para in template_paras:
                # Check if this paragraph contains a placeholder that corresponds to a list
                list_key_in_para = None
                for key, value in item_data.items():
                    if f"{{{{{key}}}}}" in template_para.text and isinstance(value, list):
                        list_key_in_para = key
                        break
                
                if list_key_in_para:
                    # This paragraph is a template for a list expansion
                    list_items = item_data[list_key_in_para]
                    placeholder = f"{{{{{list_key_in_para}}}}}"

                    # Use the first item to populate the initial paragraph
                    first_item_para = doc.add_paragraph()
                    first_item_para._p = deepcopy(template_para._p)
                    
                    if list_items:
                        _replace_placeholder_in_runs(first_item_para, placeholder, str(list_items[0]))
                    else:
                        _replace_placeholder_in_runs(first_item_para, placeholder, "") # Handle empty list
                    
                    # Also replace any other non-list placeholders in this paragraph
                    for key, value in item_data.items():
                        if key != list_key_in_para and not isinstance(value, list):
                            _replace_placeholder_in_runs(first_item_para, f"{{{{{key}}}}}", str(value))
                    
                    generated_paragraph_elements.append(first_item_para._p)

                    # Now, add subsequent list items as new paragraphs
                    if list_items and len(list_items) > 1:
                        original_font_props = None
                        for run in first_item_para.runs:
                            # A bit of a heuristic to find the run to copy style from
                            if str(list_items[0]) in run.text: 
                                original_font_props = {
                                    'bold': run.bold, 'italic': run.italic, 'underline': run.underline,
                                    'name': run.font.name, 'size': run.font.size, 'color': run.font.color.rgb
                                }
                                break
                        
                        for item in list_items[1:]:
                            new_list_item_para = doc.add_paragraph()
                            new_list_item_para.style = first_item_para.style
                            new_list_item_para._p.insert(0, deepcopy(first_item_para._p.pPr) if first_item_para._p.pPr is not None else OxmlElement('w:pPr'))
                            set_paragraph_text_with_formatting(new_list_item_para, str(item), original_font_props)
                            generated_paragraph_elements.append(new_list_item_para._p)
                else:
                    # This paragraph does not contain a list placeholder, process as normal
                    new_p = doc.add_paragraph()
                    new_p._p = deepcopy(template_para._p)
                    for key, value in item_data.items():
                        if not isinstance(value, list):
                            _replace_placeholder_in_runs(new_p, f"{{{{{key}}}}}", str(value))
                    generated_paragraph_elements.append(new_p._p)
        
        # Insert all generated paragraphs after the anchor element
        # If anchor_element is None, it means the opening_para was the first in the document
        # In this case, prepend to the document body.
        if anchor_element is None:
            # Need to find the document body element
            body = doc._body
            for p_elem in reversed(generated_paragraph_elements): # Insert in reverse to maintain order
                body.insert(0, p_elem)
        else:
            current_anchor_for_insertion = anchor_element
            for p_elem in generated_paragraph_elements:
                current_anchor_for_insertion.addnext(p_elem)
                current_anchor_for_insertion = p_elem
            
        # Rebuild paragraphs list and restart scan
        all_paragraphs = list(doc.paragraphs)
        i = 0
        continue # Continue outer loop from the start after modifying doc.paragraphs

    # Now, expand list placeholders (KEY_ACHIEVEMENTS, TECHNICAL_STACK, DESCRIPTION)
    # This must be done after repeating blocks are processed, as these can contain lists too
    
    paragraphs_to_process = list(doc.paragraphs) # Get a fresh list
    for para in paragraphs_to_process:
        for key in data:
            if isinstance(data[key], list) and (
               (key == 'KEY_ACHIEVEMENTS' and f"{{{{{key}}}}}" in para.text) or
               (key == 'TECHNICAL_STACK' and f"{{{{{key}}}}}" in para.text)
            ):
                placeholder = f"{{{{{key}}}}}"
                if placeholder in para.text:
                    list_items = data[key]
                    if list_items:
                        # Get original paragraph style and properties
                        original_style = para.style
                        original_pPr_element = deepcopy(para._p.pPr) if para._p.pPr is not None else OxmlElement('w:pPr')
                        original_font_props = None
                        # Find the font properties of the run containing the placeholder
                        for run in para.runs:
                            if placeholder in run.text:
                                original_font_props = {
                                    'bold': run.bold, 'italic': run.italic, 'underline': run.underline,
                                    'name': run.font.name, 'size': run.font.size, 'color': run.font.color.rgb
                                }
                                break

                        # Replace the placeholder in the existing paragraph with the first item
                        _replace_placeholder_in_runs(para, placeholder, str(list_items[0]))
                        
                        para.style = original_style # Reapply original style
                        
                        # Clear existing pPr if any and re-append the original pPr element
                        existing_pPr = para._p.pPr
                        if existing_pPr is not None:
                            para._p.remove(existing_pPr)
                        para._p.insert(0, original_pPr_element)
                        
                        # Insert remaining items as new paragraphs with same style and properties
                        current_anchor_elem = para._p
                        for item in list_items[1:]:
                            new_para = doc.add_paragraph()
                            new_para.style = original_style # Copy style
                            
                            # Copy pPr element
                            new_para._p.insert(0, deepcopy(original_pPr_element))
                            
                            set_paragraph_text_with_formatting(new_para, str(item), original_font_props)

                            current_anchor_elem.addnext(new_para._p)
                            current_anchor_elem = new_para._p
                    else:
                        # If list is empty, remove placeholder
                        _replace_placeholder_in_runs(para, placeholder, "") # Replace with empty string
    
    # Finally, fill any remaining simple top-level placeholders
    # Need to re-list paragraphs as they might have changed due to block processing
    for para in doc.paragraphs:
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if not isinstance(value, (list, dict)) and placeholder in para.text: # Only fill simple placeholders
                _replace_placeholder_in_runs(para, placeholder, str(value))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, value in data.items():
                        placeholder = f"{{{{{key}}}}}"
                        if not isinstance(value, (list, dict)) and placeholder in para.text: # Only fill simple placeholders
                            _replace_placeholder_in_runs(para, placeholder, str(value))
    
    return doc


def main():
    parser = argparse.ArgumentParser(
        description='Fill CV template with provided data',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Using JSON data file
  python cv_generator.py --template cv_template.txt --output cv_filled.txt --data data.json
  
  # Using command-line arguments
  python cv_generator.py --template cv_generator.py --template cv_filled.txt --name "John Doe" --email "john@example.com"
  
  # Mix JSON file and command-line arguments (command-line takes precedence)
  python cv_generator.py --template cv_template.txt --output cv_filled.txt --data data.json --name "Jane Smith"
        """
    )
    
    parser.add_argument('--template', required=True, help='Path to the CV template file')
    parser.add_argument('--output', required=True, help='Path to save the filled CV')
    parser.add_argument('--data', help='Path to JSON file with replacement data')
    parser.add_argument('--pdf', action='store_true', help='Also generate a PDF file from the DOCX')
    parser.add_argument('--role', help='Specify the role for organizing output files')
    parser.add_argument('--company', help='Specify the company name for organizing output files')
    
    args = parser.parse_args()
    
    # Load template
    template_doc = load_template(args.template)
    
    # Initialize data dictionary
    data = {}
    
    # Load from JSON file if provided
    if args.data:
        data = load_data_from_json(args.data)
    
    # Fill the template
    filled_doc = fill_docx_template(template_doc, data)

    # Remove any trailing empty paragraphs to prevent extra blank pages
    remove_trailing_empty_paragraphs(filled_doc)
    
    # Output result
    try:
        output_path_base = Path(args.output)
        
        if args.role and args.company:
            # Save to a dedicated "CVs" folder in the user's Documents directory as a workaround for macOS sandboxing
            base_cv_dir = Path.home() / "Documents" / "CVs"
            role_dir = base_cv_dir / args.role.replace(" ", "-")
            company_dir = role_dir / args.company.replace(" ", "-")
            
            output_dir = company_dir
            output_dir.mkdir(parents=True, exist_ok=True, mode=0o755)
            
            # Construct the new filename based on role
            formatted_role = args.role.replace(" ", "-")
            new_file_name = f"Jon-Hidalgo-CV-{formatted_role}{output_path_base.suffix}"
            output_path = output_dir / new_file_name
        elif args.role:
            # If only --role is provided, use the same CVs directory structure
            base_cv_dir = Path.home() / "Documents" / "CVs"
            formatted_role = args.role.replace(" ", "-")
            new_file_name = f"Jon-Hidalgo-CV-{formatted_role}{output_path_base.suffix}"
            output_path = base_cv_dir / new_file_name
            output_path.parent.mkdir(parents=True, exist_ok=True) # Ensure parent directory exists
        else:
            # If --role or --company not provided, use original output path logic
            output_dir = output_path_base.parent
            output_dir.mkdir(parents=True, exist_ok=True)
            output_path = output_path_base
        
        # Determine the final output paths for DOCX and PDF
        final_docx_output_path = output_path
        
        # Save the DOCX file
        filled_doc.save(str(final_docx_output_path))
        print(f"✓ CV filled successfully: {final_docx_output_path}")

        if args.pdf:
            try:
                final_pdf_output_path = output_path.with_suffix('.pdf')
                
                # Construct the soffice command based on the successful manual execution
                soffice_command = [
                    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
                    "--headless",
                    "--nologo",
                    "--nolockcheck",
                    "--norestore",
                    "--nodefault",
                    "--convert-to", "pdf",
                    str(final_docx_output_path), # Input from the final DOCX path
                    "--outdir", str(final_pdf_output_path.parent) # Output to the final directory
                ]

                # Run soffice command
                result = subprocess.run(soffice_command, capture_output=True, text=True, check=True)
                
                # Verify that the PDF was created
                if final_pdf_output_path.exists():
                    print(f"✓ PDF generated successfully: {final_pdf_output_path}")
                else:
                    # Look for the generated file as soffice might name it differently
                    generated_pdf_name = final_docx_output_path.with_suffix('.pdf').name
                    possible_pdf_path = final_pdf_output_path.parent / generated_pdf_name
                    if possible_pdf_path.exists():
                         print(f"✓ PDF generated successfully: {possible_pdf_path}")
                    else:
                        print("Error: PDF conversion reported success, but the output file was not found.", file=sys.stderr)
                        if result.stderr:
                             print(f"soffice stderr: {result.stderr}", file=sys.stderr)
                        sys.exit(1)
            
            except FileNotFoundError:
                print("Error: 'soffice' (LibreOffice executable) not found at /Applications/LibreOffice.app/Contents/MacOS/soffice.", file=sys.stderr)
                print("Please ensure LibreOffice is installed in the standard Applications directory.", file=sys.stderr)
                sys.exit(1)
            except subprocess.CalledProcessError as e:
                print(f"Error generating PDF with LibreOffice: {e}", file=sys.stderr)
                print(f"soffice stdout: {e.stdout}", file=sys.stderr)
                print(f"soffice stderr: {e.stderr}", file=sys.stderr)
                print("Please ensure LibreOffice is correctly installed and accessible.", file=sys.stderr)
                sys.exit(1)
            except Exception as e:
                print(f"Error generating PDF: {e}", file=sys.stderr)
                sys.exit(1)

    except Exception as e:
        print(f"Error writing output: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()