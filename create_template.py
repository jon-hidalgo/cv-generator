from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER

doc = Document()

# Set default font and line spacing for the document
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
font.color.rgb = RGBColor(0, 0, 0)
style.paragraph_format.line_spacing = 1.0
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

# Set narrow margins
sections = doc.sections
for section in sections:
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Pt(36)


def add_right_tab_stop(paragraph, section):
    usable_width = section.page_width - section.left_margin - section.right_margin
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        usable_width,
        alignment=WD_TAB_ALIGNMENT.RIGHT,
        leader=WD_TAB_LEADER.SPACES,
    )

# Title - Name
name = doc.add_paragraph()
name_run = name.add_run('{{NAME}}')
name_run.font.name = 'Times New Roman'
name_run.font.size = Pt(24)
name_run.font.bold = True
name_run.font.color.rgb = RGBColor(0, 0, 0)
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
name.paragraph_format.space_after = Pt(0)
name.paragraph_format.line_spacing = 1.0

# Subtitle - Title
subtitle = doc.add_paragraph('{{TITLE}}')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(6)
subtitle.paragraph_format.line_spacing = 1.0
for run in subtitle.runs:
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)

# Contact info
contact = doc.add_paragraph('{{LOCATION}} • {{PHONE}} • {{EMAIL}} • {{LINKEDIN}}')
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.paragraph_format.space_after = Pt(12)
contact.paragraph_format.line_spacing = 1.0
for run in contact.runs:
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)

# Horizontal line
hr1 = doc.add_paragraph()
hr1_format = hr1.paragraph_format
hr1_format.space_before = Pt(0)
hr1_format.space_after = Pt(6)

# Professional Summary
summary = doc.add_paragraph('{{PROFESSIONAL_SUMMARY}}')
summary.paragraph_format.space_after = Pt(12)
summary.paragraph_format.line_spacing = 1.0
for run in summary.runs:
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)

# Horizontal line
hr2 = doc.add_paragraph()
hr2.paragraph_format.space_after = Pt(6)

# Key Achievements Section
achievements_header = doc.add_heading('KEY ACHIEVEMENTS', level=1)
achievements_header.paragraph_format.space_before = Pt(6)
achievements_header.paragraph_format.space_after = Pt(6)
achievements_header.paragraph_format.line_spacing = 1.15
for run in achievements_header.runs:
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)

achievements = doc.add_paragraph('{{KEY_ACHIEVEMENTS}}', style='List Bullet')
achievements.paragraph_format.space_after = Pt(12)
achievements.paragraph_format.line_spacing = 1.15
for run in achievements.runs:
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)

# Technical Stack Section
tech_header = doc.add_heading('TECHNICAL STACK', level=1)
tech_header.paragraph_format.space_after = Pt(6)
tech_header.paragraph_format.line_spacing = 1.15
for run in tech_header.runs:
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)

tech = doc.add_paragraph('{{TECHNICAL_STACK}}', style='List Bullet')
tech.paragraph_format.space_after = Pt(12)
tech.paragraph_format.line_spacing = 1.15
for run in tech.runs:
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)

# Professional Experience Section
exp_header = doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)
exp_header.paragraph_format.space_after = Pt(6)
exp_header.paragraph_format.line_spacing = 1.15
for run in exp_header.runs:
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)

exp_start = doc.add_paragraph('{{#EXPERIENCE}}')
exp_start.paragraph_format.line_spacing = 1.15

exp_company = doc.add_paragraph('{{COMPANY}}\t{{JOB_LOCATION}}')
exp_company.paragraph_format.line_spacing = 1.15
add_right_tab_stop(exp_company, sections[0])
for run in exp_company.runs:
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)

exp_position = doc.add_paragraph('{{POSITION}}\t{{TIME}}')
exp_position.paragraph_format.line_spacing = 1.15
add_right_tab_stop(exp_position, sections[0])
for run in exp_position.runs:
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)

exp_description = doc.add_paragraph('{{DESCRIPTION}}', style='List Bullet')
exp_description.paragraph_format.space_after = Pt(12)
exp_description.paragraph_format.line_spacing = 1.15
for run in exp_description.runs:
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)

exp_end = doc.add_paragraph('{{/EXPERIENCE}}')
exp_end.paragraph_format.line_spacing = 1.15

# Leadership Section
leadership_header = doc.add_heading('LEADERSHIP & INTERNATIONAL PROJECTS', level=1)
leadership_header.paragraph_format.space_after = Pt(6)
leadership_header.paragraph_format.line_spacing = 1.0
for run in leadership_header.runs:
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)

projects_start = doc.add_paragraph('{{#PROJECTS}}')
projects_start.paragraph_format.line_spacing = 1.0

projects_title = doc.add_paragraph('{{ORG}} | {{POSITION}}\t{{TIME}}')
projects_title.paragraph_format.line_spacing = 1.0
add_right_tab_stop(projects_title, sections[0])
for run in projects_title.runs:
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)

projects_description = doc.add_paragraph('{{DESCRIPTION}}', style='List Bullet')
projects_description.paragraph_format.space_after = Pt(12)
projects_description.paragraph_format.line_spacing = 1.0
for run in projects_description.runs:
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)

projects_end = doc.add_paragraph('{{/PROJECTS}}')
projects_end.paragraph_format.line_spacing = 1.0

# Education Section
edu_header = doc.add_heading('EDUCATION', level=1)
edu_header.paragraph_format.space_after = Pt(6)
edu_header.paragraph_format.line_spacing = 1.0
for run in edu_header.runs:
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)

education = doc.add_paragraph('{{EDUCATION}}')
education.paragraph_format.space_after = Pt(12)
education.paragraph_format.line_spacing = 1.0
for run in education.runs:
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)

# Languages Section
lang_header = doc.add_heading('LANGUAGES', level=1)
lang_header.paragraph_format.space_after = Pt(6)
lang_header.paragraph_format.line_spacing = 1.0
for run in lang_header.runs:
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)

languages = doc.add_paragraph('{{LANGUAGES}}')
languages.paragraph_format.line_spacing = 1.0
for run in languages.runs:
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)

# Save
doc.save('example_template.docx')
print("✓ Created example_template.docx matching your CV format")
